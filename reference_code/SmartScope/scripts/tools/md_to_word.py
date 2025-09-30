#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word 转换脚本
将 Markdown 文件转换为 Word 文档
"""

import re
import sys
import os
from pathlib import Path

def install_required_packages():
    """安装必要的包"""
    import subprocess
    
    packages = ['python-docx', 'markdown']
    
    for package in packages:
        try:
            __import__(package.replace('-', '_'))
            print(f"✓ {package} 已安装")
        except ImportError:
            print(f"正在安装 {package}...")
            try:
                subprocess.check_call([sys.executable, '-m', 'pip', 'install', package])
                print(f"✓ {package} 安装成功")
            except subprocess.CalledProcessError:
                print(f"✗ {package} 安装失败")
                return False
    return True

def markdown_to_word(md_file, output_file=None):
    """将 Markdown 文件转换为 Word 文档"""
    
    # 确保安装了必要的包
    if not install_required_packages():
        print("无法安装必要的依赖包")
        return False
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        import markdown
    except ImportError as e:
        print(f"导入模块失败: {e}")
        return False
    
    # 读取 Markdown 文件
    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
    except FileNotFoundError:
        print(f"文件不存在: {md_file}")
        return False
    except Exception as e:
        print(f"读取文件失败: {e}")
        return False
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    
    # 处理 Markdown 内容
    lines = md_content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # 空行
            doc.add_paragraph()
            continue
        
        # 处理标题
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title_text = line.lstrip('#').strip()
            
            if level == 1:
                heading = doc.add_heading(title_text, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(title_text, level=min(level, 9))
            continue
        
        # 处理代码块
        if line.startswith('```'):
            # 跳过代码块标记行
            continue
        
        # 处理列表
        if line.startswith('- ') or line.startswith('* '):
            list_text = line[2:].strip()
            # 移除 Markdown 格式
            list_text = re.sub(r'\*\*(.*?)\*\*', r'\1', list_text)  # 粗体
            list_text = re.sub(r'\*(.*?)\*', r'\1', list_text)      # 斜体
            list_text = re.sub(r'`(.*?)`', r'\1', list_text)        # 代码
            doc.add_paragraph(list_text, style='List Bullet')
            continue
        
        # 处理数字列表
        if re.match(r'^\d+\.', line):
            list_text = re.sub(r'^\d+\.\s*', '', line)
            # 移除 Markdown 格式
            list_text = re.sub(r'\*\*(.*?)\*\*', r'\1', list_text)
            list_text = re.sub(r'\*(.*?)\*', r'\1', list_text)
            list_text = re.sub(r'`(.*?)`', r'\1', list_text)
            doc.add_paragraph(list_text, style='List Number')
            continue
        
        # 处理普通段落
        if line:
            # 移除 Markdown 格式
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # 粗体
            text = re.sub(r'\*(.*?)\*', r'\1', text)      # 斜体
            text = re.sub(r'`(.*?)`', r'\1', text)        # 内联代码
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # 链接
            
            # 跳过 Mermaid 图表和其他特殊内容
            if 'mermaid' in text.lower() or text.startswith('graph ') or text.startswith('```'):
                continue
                
            doc.add_paragraph(text)
    
    # 确定输出文件名
    if output_file is None:
        md_path = Path(md_file)
        output_file = md_path.with_suffix('.docx')
    
    # 保存文档
    try:
        doc.save(output_file)
        print(f"✓ Word 文档已保存: {output_file}")
        return True
    except Exception as e:
        print(f"保存文档失败: {e}")
        return False

def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("用法: python md_to_word.py <markdown_file> [output_file]")
        print("示例: python md_to_word.py 软件设计说明书.md")
        return
    
    md_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    if not os.path.exists(md_file):
        print(f"文件不存在: {md_file}")
        return
    
    print(f"正在转换 {md_file} 为 Word 文档...")
    
    if markdown_to_word(md_file, output_file):
        print("转换完成!")
    else:
        print("转换失败!")

if __name__ == "__main__":
    main()
